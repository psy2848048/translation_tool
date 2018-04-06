from app import app, db, mail
from sqlalchemy import Table, MetaData, text, and_
import traceback
from datetime import datetime
import csv
import io
import os


def select_doc(did):
    conn = db.engine.connect()

    res = conn.execute(text("""SELECT os.id as sentence_id, d.origin_lang, d.trans_lang, os.text as origin_text
                                        , IF(ts.text is not NULL, ts.text, '') as trans_text
                                        , IF(ts.status is not NULL, ts.status, 0) as trans_status
                                        , IF(ts.type is not NULL, ts.type, 0) as trans_type
                                        , IF(comment_cnt is not NULL, comment_cnt, 0) as comment_cnt
                                  FROM `marocat v1.1`.doc_origin_sentences os JOIN docs d ON d.id = os.doc_id
                                  											  LEFT JOIN doc_trans_sentences ts ON ts.origin_id = os.id AND ts.is_deleted = FALSE
																			  LEFT JOIN ( SELECT origin_id, COUNT(*) as comment_cnt FROM trans_comments 
								                                                          WHERE is_deleted = FALSE GROUP BY origin_id ) tc ON tc.origin_id = os.id
                                  WHERE os.doc_id = :did AND os.is_deleted = FALSE;"""), did=did)

    doc_sentences = [dict(r) for r in res]
    return doc_sentences


def export_doc(output_type, did):
    conn = db.engine.connect()
    meta = MetaData(bind=db.engine)

    tm = Table('translation_memory', meta, autoload=True)
    ut = Table('user_tmlist', meta, autoload=True)

    #: 문장들 꺼내기
    res = conn.execute(text("""SELECT os.id as osid
                                      , d.project_id as pid, d.id as did, d.title
                                      , origin_lang, trans_lang
                                      , os.text as origin_text
                                      , IF(ts.text is not NULL, ts.text, '') as trans_text
                              FROM `marocat v1.1`.doc_origin_sentences os JOIN docs d ON d.id = os.doc_id
                                                                          LEFT JOIN doc_trans_sentences ts ON ts.origin_id = os.id AND ts.is_deleted = FALSE
                              WHERE os.doc_id = :did AND os.is_deleted = FALSE;"""), did=did).fetchall()

    doc_title = res[0]['title']
    udate = str(datetime.utcnow().strftime('%Y%m%d%H%M%S'))
    file_title = 'mycattool-' + str(res[0]['pid']) + '-' + str(res[0]['did']) + '-' + udate + '.' + output_type

    file = write_file_in_requested_format(output_type, doc_title, res)

    #: 관리자 계정의 문장저장소에 결과 넣기(일단, 무조건 ciceron@ciceron.me에 추가)
    try:
        for r in res:
            res2 = conn.execute(tm.insert()
                                , origin_lang=r['origin_lang'], trans_lang=r['trans_lang']
                                , origin_text=r['origin_text'], trans_text=r['trans_text'])
            tid = res2.lastrowid
            conn.execute(ut.insert(), user_id=7, doc_id=did, tm_id=tid)
    except:
        print('Wrong! (insert tm)')
        traceback.print_exc()
        pass

    return (file, file_title), True


def write_file_in_requested_format(output_type, doc_title, res):
    #: CSV 파일로 출력하기
    if output_type == 'csv':
        output = io.StringIO()
        writer = csv.writer(output, quoting=csv.QUOTE_NONNUMERIC)
        writer.writerows(res)
        file = output.getvalue().encode('utf-8')

    #: TXT 또는 Markdown 파일로 출력하기
    elif output_type == 'txt' or output_type == 'md':
        with open('output.txt', 'w') as f:
            f.write('제목: ' + doc_title + '\n')
            f.write('원문언어: ' + res[0]['origin_lang'].upper() + '\n')
            f.write('번역언어: ' + res[0]['trans_lang'].upper() + '\n\n')

            for r in res:
                f.write(r['origin_text'] + '\n')

            f.write('\n\n')

            for r in res:
                f.write(r['trans_text'] + '\n')

            f.write('\n\n')
            f.write('-' * 2 * len(res[0]['trans_text']) + '\n\n')

            for r in res:
                a = '{}-{}: '.format(r['osid'], r['origin_lang'].upper())
                f.write(a)
                f.write(r['origin_text'] + '\n')

                b = '{}-{}: '.format(r['osid'], r['trans_lang'].upper())
                f.write(b)
                f.write(r['trans_text'] + '\n\n')

        file = open('output.txt', 'rb').read()
        os.remove('output.txt')

    #: DOCX 파일로 출력하기
    elif output_type == 'docx':
        from docx import Document
        doc = Document()
        doc.add_heading('{}'.format(doc_title), 0)

        doc.add_paragraph('원문언어: ' + res[0]['origin_lang'].upper())
        doc.add_paragraph('번역언어: ' + res[0]['trans_lang'].upper() + '\n')

        doc.add_heading('원문', level=1)
        for r in res:
            doc.add_paragraph(r['origin_text'])

        doc.add_page_break()

        doc.add_heading('번역문', level=1)
        for r in res:
            doc.add_paragraph(r['trans_text'])

        doc.add_page_break()

        for r in res:
            a = '{}-{}: '.format(r['osid'], r['origin_lang'].upper())
            doc.add_paragraph(a + r['origin_text'])

            b = '{}-{}: '.format(r['osid'], r['trans_lang'].upper())
            doc.add_paragraph(b + r['trans_text'])

        doc.save('output.docx')
        file = open('output.docx', 'rb').read()
        os.remove('output.docx')

    else:
        file = None

    return file


def insert_or_update_trans(sid, trans_text, trans_type):
    conn = db.engine.connect()
    trans = conn.begin()

    try:
        res = conn.execute(text("""INSERT INTO `marocat v1.1`.doc_trans_sentences
                                   SET origin_id = :oid, text = :trans_text, type = :trans_type
                                   ON DUPLICATE KEY UPDATE origin_id = :oid, text = :trans_text, type = :trans_type, update_time = CURRENT_TIMESTAMP;""")
                           , oid=sid, trans_text=trans_text, trans_type=trans_type)
        if res.rowcount not in [1, 2]:
            trans.rollback()
            return False

        trans.commit()
        return True
    except:
        traceback.print_exc()
        trans.rollback()
        return False


def update_sentence_status(sid, status):
    conn = db.engine.connect()
    trans = conn.begin()
    meta = MetaData(bind=db.engine)
    ts = Table('doc_trans_sentences', meta, autoload=True)

    try:
        res = conn.execute(ts.update(ts.c.origin_id == sid), status=status, update_time=datetime.utcnow())
        if res.rowcount != 1:
            trans.rollback()
            return False

        trans.commit()

        #: 100%라면 관리자에게 완료 이메일 보내기
        check_doc_comeplte(sid)

        return True
    except:
        traceback.print_exc()
        trans.rollback()
        return False


def check_doc_comeplte(sid):
    conn = db.engine.connect()
    meta = MetaData(bind=db.engine)
    dm = Table('doc_members', meta, autoload=True)

    #: 번역 상태 100%인지 확인
    res1 = conn.execute(text(
        """SELECT d.project_id as pid, p.name as project_name
                , d.id as did, d.title as doc_title, d.origin_lang, d.trans_lang
                , IF(CAST(FLOOR(SUM(ts.status) / COUNT(*) * 100) AS CHAR) is not NULL
                , CAST(FLOOR(SUM(ts.status) / COUNT(*) * 100) AS CHAR), 0) as progress_percent
           FROM `marocat v1.1`.docs d JOIN projects p ON d.project_id = p.id
           LEFT JOIN (doc_origin_sentences os, doc_trans_sentences ts) ON (os.doc_id = d.id AND os.id = ts.origin_id)
           WHERE d.id = (SELECT doc_id FROM doc_origin_sentences WHERE id = 4160)
           AND ts.is_deleted = FALSE AND os.is_deleted = FALSE"""), sid=sid).fetchone()

    #: 100% 아니라면 패스
    progress_percent = int(res1['progress_percent'])
    print(111, progress_percent)
    if int(progress_percent) == 100:
        #: 관리자 이메일 꺼내기
        res2 = conn.execute(text(
            """SELECT email
               FROM project_members pm 
               JOIN (users u, docs d) ON (u.id = pm.user_id AND d.project_id=pm.project_id)
               JOIN (SELECT doc_id as did FROM doc_origin_sentences WHERE id = :sid) t1 ON t1.did = d.id
               WHERE is_admin = True"""), sid=sid).fetchall()
        mail_to = [r['email'] for r in res2]
        print(222, mail_to)

        project_name = res1['project_name']
        doc_title = res1['doc_title']
        content = '`{}` 프로젝트의 `{}` 문서의 번역 작업이 완료됐습니다.'.format(project_name, doc_title)

        mail.send_mail_directly(mail_from=app.config['MANAGER_MAIL'], mail_to=mail_to
                                , subject='번역이 완료되었습니다', content=content)
        return True
    else:
        return False


def select_trans_comments(sid):
    conn = db.engine.connect()
    results = conn.execute(text("""SELECT c.id as comment_id, doc_id, origin_id as sentence_id, user_id, u.name, text as comment, c.create_time
                                   FROM `marocat v1.1`.trans_comments c JOIN users u ON u.id = c.user_id
                                   WHERE origin_id = :sid AND c.is_deleted = FALSE AND u.is_deleted = FALSE
                                   ORDER BY c.create_time;"""), sid=sid).fetchall()
    comments = [dict(res) for res in results]
    return comments


def insert_trans_comment(uid, did, sid, comment):
    conn = db.engine.connect()
    trans = conn.begin()
    meta = MetaData(bind=db.engine)
    c = Table('trans_comments', meta, autoload=True)

    try:
        res = conn.execute(c.insert(), user_id=uid, doc_id=did, origin_id=sid, text=comment)
        if res.rowcount != 1:
            trans.rollback()
            return 0

        trans.commit()
        return True
    except:
        traceback.print_exc()
        trans.rollback()
        return False


def delete_trans_comment(cid):
    conn = db.engine.connect()
    trans = conn.begin()
    meta = MetaData(bind=db.engine)
    sc = Table('trans_comments', meta, autoload=True)

    try:
        res = conn.execute(sc.update(sc.c.id == cid), is_deleted=True, update_time=datetime.utcnow())
        if res.rowcount != 1:
            trans.rollback()
            return 0

        trans.commit()
        return True
    except:
        traceback.print_exc()
        trans.rollback()
        return False


def select_doc_comments(did, page, rows):
    conn = db.engine.connect()

    res = conn.execute(text("""SELECT count(*) FROM `marocat v1.1`.trans_comments WHERE doc_id = :did AND is_deleted = FALSE;""")
                       , did=did).fetchone()
    total_cnt = res[0]

    results = conn.execute(text("""SELECT c.id as comment_id, doc_id, origin_id as sentence_id, user_id, u.name, text as comment, c.create_time
                                   FROM `marocat v1.1`.trans_comments c JOIN users u ON u.id = c.user_id
                                   WHERE doc_id = :did AND c.is_deleted = FALSE AND u.is_deleted = FALSE
                                   ORDER BY c.create_time
                                  LIMIT :row_count OFFSET :offset;""")
                           , did=did, row_count=rows, offset=rows * (page - 1)).fetchall()

    comments = [dict(res) for res in results]
    return comments, total_cnt


def insert_doc_comment(uid, did, comment):
    conn = db.engine.connect()
    trans = conn.begin()
    meta = MetaData(bind=db.engine)
    c = Table('trans_comments', meta, autoload=True)

    try:
        res = conn.execute(c.insert(), user_id=uid, doc_id=did, text=comment)
        if res.rowcount != 1:
            trans.rollback()
            return 0

        trans.commit()
        return True
    except:
        traceback.print_exc()
        trans.rollback()
        return False


def delete_doc_comment(cid):
    conn = db.engine.connect()
    trans = conn.begin()
    meta = MetaData(bind=db.engine)
    sc = Table('trans_comments', meta, autoload=True)

    try:
        res = conn.execute(sc.update(sc.c.id == cid), is_deleted=True, update_time=datetime.utcnow())
        if res.rowcount != 1:
            trans.rollback()
            return 0

        trans.commit()
        return True
    except:
        traceback.print_exc()
        trans.rollback()
        return False

